import streamlit as st
import faiss
import os
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.llms import HuggingFaceEndpoint

#from langchain_huggingface import HuggingFaceEndpoint

#from secret_api_keys import huggingface_api_key  # Set the Hugging Face Hub API token as an environment variable
#os.environ['HUGGINGFACEHUB_API_TOKEN'] = huggingface_api_key
HUGGINGFACEHUB_API_TOKEN='hf_yourtokenheredvQR'


def process_input(input_type, input_data):
    """Processes different input types and returns a vectorstore."""
    loader = None
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for PDF")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data  # Input is already a text string
        else:
            raise ValueError("Expected a string for 'Text' input type.")
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for DOCX")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, UploadedFile):
            text = str(input_data.read().decode('utf-8'))
        else:
            raise ValueError("Invalid input data for TXT")
        documents = text
    else:
        raise ValueError("Unsupported input type")

    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [ str(doc.page_content) for doc in texts ]  # Access page_content from each Document 
    else:
        texts = text_splitter.split_text(documents)

    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}

    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )
    # Create FAISS index
    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)
    # Create FAISS vector store with the embedding function
    vector_store = FAISS(
        embedding_function=hf_embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)  # Add documents to the vector store
    return vector_store

def answer_question(vectorstore, query):
    """Answers a question based on the provided vectorstore."""
    llm = HuggingFaceEndpoint(repo_id= 'meta-llama/Meta-Llama-3-8B-Instruct', 
                              token = huggingface_api_key, temperature= 0.6)
    qa = RetrievalQA.from_chain_type(llm=llm, retriever=vectorstore.as_retriever())

    answer = qa({"query": query})
    return answer

def main():
    st.title("RAG Q&A App")
    input_type = st.selectbox("Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])
    if input_type == "Link":
        number_input = st.number_input(min_value=1, max_value=20, step=1, label = "Enter the number of Links")
        input_data = []
        for i in range(number_input):
            url = st.sidebar.text_input(f"URL {i+1}")
            input_data.append(url)
    elif input_type == "Text":
        input_data = st.text_input("Enter the text")
    elif input_type == 'PDF':
        input_data = st.file_uploader("Upload a PDF file", type=["pdf"])
    elif input_type == 'TXT':
        input_data = st.file_uploader("Upload a text file", type=['txt'])
    elif input_type == 'DOCX':
        input_data = st.file_uploader("Upload a DOCX file", type=[ 'docx', 'doc'])
    if st.button("Proceed"):
        # st.write(process_input(input_type, input_data))
        vectorstore = process_input(input_type, input_data)
        st.session_state["vectorstore"] = vectorstore
    if "vectorstore" in st.session_state:
        query = st.text_input("Ask your question")
        if st.button("Submit"):
            answer = answer_question(st.session_state["vectorstore"], query)
            st.write(answer)

if __name__ == "__main__":
    main()